#!/usr/bin/env python3
"""
Word文書を一括でプレーンテキストに変換するスクリプト

使い方:
    python word_to_text.py [フォルダパス]

例:
    python word_to_text.py C:\Users\Documents\WordFiles
    python word_to_text.py .  # 現在のフォルダ

機能:
- .docxファイルをプレーンテキストに変換
- サブフォルダも再帰的に検索
- 元のファイル名.txtで保存
- 進捗表示とエラーハンドリング
"""

import os
import sys
from pathlib import Path
from docx import Document


def convert_docx_to_text(docx_path, output_path=None):
    """
    .docxファイルをプレーンテキストに変換

    Args:
        docx_path: Wordファイルのパス
        output_path: 出力先パス（Noneの場合は元のファイル名.txt）

    Returns:
        bool: 成功した場合True
    """
    try:
        # Wordファイルを読み込み
        doc = Document(docx_path)

        # 全段落のテキストを抽出
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)

        # テーブル内のテキストも抽出
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        # 出力パスを決定
        if output_path is None:
            output_path = str(Path(docx_path).with_suffix('.txt'))

        # テキストファイルに保存
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(full_text))

        return True

    except Exception as e:
        print(f'❌ エラー: {docx_path}')
        print(f'   {str(e)}')
        return False


def find_word_files(directory):
    """
    指定ディレクトリ内の全.docxファイルを再帰的に検索

    Args:
        directory: 検索するディレクトリ

    Returns:
        list: .docxファイルのパスリスト
    """
    word_files = []
    directory_path = Path(directory)

    # .docxファイルを再帰的に検索
    for file_path in directory_path.rglob('*.docx'):
        # 一時ファイル（~$で始まる）は除外
        if not file_path.name.startswith('~$'):
            word_files.append(file_path)

    return word_files


def main():
    """メイン処理"""

    # コマンドライン引数からフォルダパスを取得
    if len(sys.argv) > 1:
        target_dir = sys.argv[1]
    else:
        target_dir = '.'

    # フォルダの存在確認
    if not os.path.exists(target_dir):
        print(f'❌ フォルダが見つかりません: {target_dir}')
        sys.exit(1)

    print(f'📁 検索中: {os.path.abspath(target_dir)}')
    print()

    # Wordファイルを検索
    word_files = find_word_files(target_dir)

    if not word_files:
        print('❌ .docxファイルが見つかりませんでした')
        sys.exit(0)

    print(f'📄 {len(word_files)}個のWordファイルが見つかりました')
    print()

    # 変換処理
    success_count = 0
    fail_count = 0

    for i, word_file in enumerate(word_files, 1):
        print(f'[{i}/{len(word_files)}] {word_file.name}', end=' ... ')

        if convert_docx_to_text(word_file):
            print('✅ 完了')
            success_count += 1
        else:
            fail_count += 1

    # 結果サマリー
    print()
    print('=' * 50)
    print(f'✅ 成功: {success_count}件')
    if fail_count > 0:
        print(f'❌ 失敗: {fail_count}件')
    print('=' * 50)


if __name__ == '__main__':
    main()
